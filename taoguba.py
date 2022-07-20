import requests
import re
import time
import random
from bs4 import BeautifulSoup
from urllib.request import urlretrieve
import docx
from PIL import Image
from docx.shared import Inches

TMP_IMG_NAME = 'tmpimg.jpg'
A4_WIDTH = 590
A4_WIDTH_INCH = 6


def get_url(url):
    '''
        给出url,获取网页回应
    '''
    user_agent = 'Mozilla/5.0 (Windows NT 6.1; WOW64)'
    cookie = 'agree=enter; tgbuser=6157693; tgbpwd=BAF37F3E832qr0tzopr318k1th; JSESSIONID=M2RiZGNlMGMtNTg3Zi00ODkyLTg1Y2UtOTZjZjYwNWQ2NzQy; Hm_lvt_cc6a63a887a7d811c92b7cc41c441837=1657117241,1657341550,1657433823,1657634226; acw_tc=0b63bb3216576360537936677e0164417a04080e0a6c98122fb68e4047cb1a; Hm_lpvt_cc6a63a887a7d811c92b7cc41c441837=1657636288'
    headers = {"User-Agent": user_agent,
               "cookie": cookie
               }
    response = requests.get(url, headers=headers)
    response = response.text
    return response


def get_one_page(soup, doc, i):
    '''
        组合目的url,发送请求，返回数据，保存数据库
    '''
    if i == 1:
        title = soup.find("b", id="b_subject").get_text()
        doc.add_heading(title, 1)
        content = soup.find("div", class_="p_coten").get_text()
        doc.add_paragraph(content)
    # 跟帖
    followList = soup.findAll("div", class_="pc_p_nr")
    for follow in followList:
        doc.add_paragraph(follow.find("span", class_='pcyclspan').get_text())
        doc.add_paragraph(follow.find("a", class_='pcyclname').get_text().strip())
        doc.add_paragraph(follow.find("p", class_='p_wz').get_text().rstrip().replace('\n', ''))
        # 判断有没有图片
        img = follow.find("img", class_='lazy')
        if img is not None:
            try:
                urlretrieve(img['data-original'], TMP_IMG_NAME)
                pic = Image.open(TMP_IMG_NAME)
                if (pic.size)[0] > A4_WIDTH:
                    doc.add_picture(TMP_IMG_NAME, width=Inches(A4_WIDTH_INCH))
                else:
                    doc.add_picture(TMP_IMG_NAME)
            except:
                print('出错啦T_T')

        minText = follow.find('div', class_='quote_container')
        if minText is not None:
            doc.add_paragraph(minText.find('a').get_text())
            doc.add_paragraph('\t\t' + minText.find('span').get_text())
            doc.add_paragraph('\t\t\t\t' + minText.find('span', class_='quote_container_time left').get_text())
        doc.add_paragraph(
            "----------------------------------------------------------------------------------------------------------------------")


def main():
    indexUrl = 'https://www.taoguba.com.cn/user/blog/moreTopic?userID=6724235'
    baseUrl = 'https://www.taoguba.com.cn/'
    soupIndex = BeautifulSoup(get_url(indexUrl), "lxml")
    indexList = soupIndex.findAll('td', class_='suh')
    flag = True
    for index in indexList:
        href = baseUrl + index.find('a')['href']
        title = index.find('a')['title']
        if title == '2022.6.23 超短高阶：指数与情绪的组合':
            flag = False
        if flag:
            continue
        soupPage = BeautifulSoup(get_url(href), "lxml")
        page = int(
            re.findall('\d+', soupPage.find('div', class_='left t_page01').findAll('span')[-1].text.split('/')[1])[
                0]) + 1
        doc = docx.Document()
        print('一共' + str(page) + "页")
        for i in range(1, page):
            print("现在第" + str(i))
            url = href[0:-1] + '{page}'
            url = url.format(page=i)
            response = get_url(url)
            soup = BeautifulSoup(response, "lxml")
            get_one_page(soup, doc, i)
            time.sleep(random.randint(1, 10))
        print("完成" + title)
        doc.save(title + '.docx')


if __name__ == "__main__":
    main()
