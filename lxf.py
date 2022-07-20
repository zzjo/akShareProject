# -*- coding:utf-8 -*-__author__ = 'zhengbiqing 460356155@qq.com'
from bs4 import BeautifulSoup
from urllib.request import urlopen, urlretrieve
import re
from collections import OrderedDict
from docx import Document
from docx.shared import Inches
import docx
import os
import time
import sys
from PIL import Image

baseUrl = 'http://www.liaoxuefeng.com'
# python
# firstPageUrl = '/wiki/001374738125095c955c1e6d8bb493182103fac9270762a000'
# git
# firstPageUrl = '/wiki/0013739516305929606dd18361248578c67b8067c8c017b000'
# javascript
firstPageUrl = '/wiki/001434446689867b27157e896e74d51a89c25cc8b43bdb3000'

INDENT = 8
TMP_IMG_NAME = 'tmpimg.jpg'
A4_WIDTH = 590
A4_WIDTH_INCH = 6
PAGE_OPEN_DELAY = 1

# 保存目录信息的有序字典
content = OrderedDict()


def writePage(url, indentNum, doc):
    '''
    保存一个html页面
    :param url: 页面地址
    :param indentNum: 目录项缩进量，形成层次性的目录结构
    :param doc: 指向保存到的word文件
    :return:
    '''
    bsObj = BeautifulSoup(urlopen(url), 'lxml')
    # 找到章节标题，并保存
    title = bsObj.find('div', id="x-content").find('h4')
    doc.add_heading(title.get_text(), int(indentNum))
    # 找到文章内容主体，得到每一个p标签，遍历每个p标签
    pageContent = bsObj.find('div', class_="x-wiki-content")
    pages = pageContent.findAll('p')

    for page in pages:
        # 对img标签，得到图片并保存到word文件
        imgs = page.findAll('img')
        for img in imgs:
            imgUrl = baseUrl + img['src'] if img['src'].startswith('/') else img['src']
            urlretrieve(imgUrl, TMP_IMG_NAME)
            pic = Image.open(TMP_IMG_NAME)
            if (pic.size)[0] > A4_WIDTH:
                doc.add_picture(TMP_IMG_NAME, width=Inches(A4_WIDTH_INCH))
            else:
                doc.add_picture(TMP_IMG_NAME)
        # 保存文字
        doc.add_paragraph(page.get_text())
    # 添加分页
    doc.add_page_break()


def writeContent(url, doc):
    '''
    得到目录信息，并保存为word文件目录，但未实现链接跳转功能，python-docx的该功能在开发中
    :param url: 目录信息业html页面地址
    :param doc: 指向保存到的word文件
    :return:
    '''
    html = urlopen(url)
    bsObj = BeautifulSoup(html, 'lxml')
    # 找到符合一定特征的目录项
    i = 0
    contentList = bsObj.findAll('li', {'style': re.compile('^margin-left.*'), 'id': not None})
    for contentItem in contentList:
        i = i + 1
        if i > 10:
            break
        contentLink = contentItem.find('a')
        # 得到目录项的缩进数，表征了目录项的层次
        indent = re.findall('\d+', contentItem['style'])
        contentStr = ' ' * INDENT * int(indent[0]) + contentLink.string
        # 避免重复的目录项
        if contentStr not in content:
            # 目录信息，包括标题、链接、缩进
            content[contentStr] = [contentLink['href'], indent[0]]

    doc.add_heading(u'目录', 0)
    for item, value in content.items():
        # print item, value
        # 保存目录到word文档
        doc.add_paragraph(item)

    doc.add_page_break()


def fileName(url):
    '''
    从html中得到word文件名
    '''
    html = urlopen(url)
    bsObj = BeautifulSoup(html, 'lxml')
    return bsObj.title.string


fileName = fileName(baseUrl + firstPageUrl) + '.doc'
document = docx.Document(fileName) if os.path.exists(fileName) else docx.Document()

writeContent(baseUrl + firstPageUrl, document)
writePage(baseUrl + firstPageUrl, '1', document)

for page in content.values():
    # 调试信息
    print(baseUrl + page[0])
    writePage(baseUrl + page[0], page[1], document)
    time.sleep(PAGE_OPEN_DELAY)

# 调试时采用，只读几个页面
'''
for i in range( 1 ):
    #调试信息
    print baseUrl + content.values()[i][0]
    writePage( baseUrl + content.values()[i][0], content.values()[i][1], document )
    time.sleep( PAGE_OPEN_DELAY )
'''

document.save(fileName)

# 删除图片临时文件
try:
    os.remove(TMP_IMG_NAME)
except:
    print('remove file fail')
